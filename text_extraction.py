import os
from docx import Document
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def extract_between_markers(doc_path):
    try:
        # Load the document
        doc = Document(doc_path)
        full_text = ""

        # Extract all text from paragraphs
        for para in doc.paragraphs:
            full_text += para.text.strip() + "\n"

        # Extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text.strip() + "\n"

        # Find text between 'Additional information' and 'Follow-Up Plan'
        start_marker = 'Additional information'
        end_marker = 'Follow-Up Plan'
        start_index = full_text.find(start_marker)
        end_index = full_text.find(end_marker)

        if start_index != -1 and end_index != -1:
            # Extract the text between the markers
            extracted_text = full_text[start_index + len(start_marker):end_index].strip()
            return extracted_text
        else:
            logging.warning(f"Markers not found in {doc_path}")
            return ''
    
    except Exception as e:
        logging.error(f"Error processing {doc_path}: {e}")
        return ''

def process_files_in_folder(folder_path, output_folder):
    # Create the output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    # Iterate over files in the folder and save their extracted text
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            doc_path = os.path.join(folder_path, filename)
            logging.info(f"Extracting text from: {filename}")
            text = extract_between_markers(doc_path)
            
            if text:
                output_file = os.path.join(output_folder, f"{filename[:-5]}_extracted.txt")
                with open(output_file, 'w') as f_out:
                    f_out.write(text)
                logging.info(f"Extracted text saved to: {output_file}")
        else:
            logging.info(f"Skipping non-docx file: {filename}")

# Set your folder path containing .docx files
folder_path = '/Users/jovanwong/Pictures/Desktop/IMH/New Prompt Engineering/New forms'

# Output folder to save the extracted text files
output_folder = '/Users/jovanwong/Pictures/Desktop/IMH/New Prompt Engineering/extracted texts'

# Extract and save text from all .docx files in the folder to individual text files
process_files_in_folder(folder_path, output_folder)

logging.info(f"Text extraction complete. Extracted files saved to {output_folder}")
